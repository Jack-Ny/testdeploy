from io import BytesIO
import json
import openpyxl
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django import template
from django.core.exceptions import ObjectDoesNotExist
from django.db import transaction
from django.http import HttpResponse
from django.shortcuts import redirect, render
from django.contrib import messages
from django.shortcuts import get_object_or_404
from .models import *
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model
from .decorators import gestion_required, gestion_user_required
from xhtml2pdf import pisa
from django.template.loader import get_template
from docx import Document
from openpyxl import Workbook
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth import update_session_auth_hash

# recuperer le modele de l'utilisation
User = get_user_model()

# Generer le pdf de generale
@login_required(login_url='/login/')
@gestion_required
def generate_pdf_generale(request):
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    context = {
        'infos_test': infos_test,
        'partenariats_dict': partenariats_dict,
        'objectifs_dict': objectifs_dict
    }
    template_path = 'service/invoice_general.html'
    template = get_template(template_path)
    html = template.render(context)
    # Créer un objet HttpResponse avec le type de contenu PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_globale.pdf"'
    # Convertir le template HTML en PDF
    pisa_status = pisa.CreatePDF(html, dest=response)
    # Si la conversion a réussi, retourner la réponse avec le PDF généré
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    return response

# Generer le word de generale
@login_required(login_url='/login/')
@gestion_required
def generate_word_generale(request):
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    doc = Document()
    doc.add_heading('SYNTHÈSE GLOBALE INFORMATIONS GÉNÉRALES', level=1)

    for info in infos_test:
        doc.add_heading(f"Projet: {info.id_projet.nom}", level=2)

        # infos organisation
        doc.add_heading(f"Informations sur l'organisation", level=3)
        doc.add_paragraph(f"Nom de l'organisation : {info.nom_org}")
        doc.add_paragraph(f"Nature de l'organisation : {info.nature_org}")
        doc.add_paragraph(f"Sigle : {info.sigle}")
        doc.add_paragraph(f"Pays d'origine  : {info.pays_origine}")
        #
        doc.add_heading(f"Adresses du siege de l'organisation", level=3)
        doc.add_paragraph(f"Région : {info.region}")
        doc.add_paragraph(f"Province : {info.province}")
        doc.add_paragraph(f"Commune : {info.commune}")
        doc.add_paragraph(f"Village/Secteur/Avenue/Rue : {info.village}")
        doc.add_paragraph(f"Boite postale : {info.boite_postale}")
        doc.add_paragraph(f"Numéro de téléphone fixe : {info.numb_fixe}")
        doc.add_paragraph(f"Numéro de téléphone mobile : {info.numb_mobile}")
        doc.add_paragraph(f"Adresse mail professionnelle : {info.adresse_mail}")
        doc.add_paragraph(f"Site web : {info.site_web}")
        #
        doc.add_heading(f"Responsable de l'organisation", level=3)
        doc.add_paragraph(f"Nom et Prénom(s) : {info.nom_complet_resp}")
        doc.add_paragraph(f"Nationalité : {info.nationalite_resp}")
        doc.add_paragraph(f"Fonction(Président,...) : {info.fonction_resp}")
        doc.add_paragraph(f"Numéro fixe : {info.numb_fixe_resp}")
        doc.add_paragraph(f"Numéro mobile : {info.numb_mobile_resp}")
        #
        doc.add_heading(f"Gouvernance interne de l'association : Tenue des rencontres statuaires des instances de l'organisation", level=3)
        doc.add_paragraph(f"Dernier renouvèlement des Instances dirigeantes : {info.renou_instance}")
        doc.add_paragraph(f"Dernière Assemblée Générale Ordinaire  : {info.assem_general}")
        doc.add_paragraph(f"Dernière session statutaire du bureau exécutif : {info.session_statut}")
        doc.add_paragraph(f"Durée du mandat du bureau exécutif : {info.mandat_bureau} ans")
        #
        doc.add_heading(f"Objectifs principaux de l'organisation", level=3)
        objectifs = objectifs_dict[info.id]
        objectifs_str = '\n'.join([o['objectifs'] for o in objectifs]) + '\n'
        doc.add_paragraph(f"Objectifs : \n {objectifs_str}")
        #
        doc.add_heading(f"Groupes cibles specifique", level=3)
        doc.add_paragraph(f"Groupes : {info.groupes_cibles}")
        #
        doc.add_heading(f"Personnel employe", level=3)
        doc.add_heading(f"Employés nationaux Contrat à Durée Indéterminée (CDI)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_nation_cdi_homme}")
        doc.add_paragraph(f"Femmes : {info.em_nation_cdi_femme}")
        doc.add_heading(f"Employés nationaux Contrat à Durée déterminée (CDD)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_nation_cdd_homme}")
        doc.add_paragraph(f"Femmes : {info.em_nation_cdd_femme}")
        doc.add_heading(f"Employés expatriés Contrat à Durée Indéterminée (CDI)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_expa_cdi_homme}")
        doc.add_paragraph(f"Femmes : {info.em_expa_cdi_femme}")
        doc.add_heading(f"Employés expatriés Contrat à Durée déterminée (CDD)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_expa_cdd_homme}")
        doc.add_paragraph(f"Femmes : {info.em_expa_cdd_femme}")
        #
        doc.add_heading(f"Bénévoles ou volontaires", level=3)
        doc.add_heading(f"Bénévoles ou volontaires Nationaux", level=4)
        doc.add_paragraph(f"Hommes : {info.benevol_nation_homme}")
        doc.add_paragraph(f"Femmes : {info.benevol_nation_femme}")
        doc.add_heading(f"Bénévoles ou volontaires Expatriés", level=4)
        doc.add_paragraph(f"Hommes : {info.benevol_expa_homme}")
        doc.add_paragraph(f"Femmes : {info.benevol_expa_femme}")
        #
        doc.add_heading(f"Personnel de l'Administration publique en détachement", level=3)
        doc.add_paragraph(f"Hommes : {info.personnel_admin_homme}")
        doc.add_paragraph(f"Femmes : {info.personnel_admin_femme}")
        #
        doc.add_heading(f"Partenariats / collaborations", level=3)
        partenariats = partenariats_dict[info.id]
        if partenariats:
            for p in partenariats:
                doc.add_paragraph(f"Nom du partenariat : {p['nom']}")
                doc.add_paragraph(f"N° de convention de partenariat / protocole d'entente : {p['numero']}")
                doc.add_paragraph(f"Date de début d'effet : {p['date_debut']}")
                doc.add_paragraph(f"Date de fin d'effet : {p['date_fin']}")
                doc.add_paragraph("----------------------------------------")
        else:
            doc.add_paragraph(f"Nom du partenariat : ")
            doc.add_paragraph(f"N° de convention de partenariat / protocole d'entente : ")
            doc.add_paragraph(f"Date de début d'effet : ")
            doc.add_paragraph(f"Date de fin d'effet : ")

        doc.add_paragraph("_________________________________________________________________________________________________________")
    output = BytesIO()
    doc.save(output)

    output.seek(0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=synthese_generales.docx'
    response.write(output.getvalue())

    return response


# Generer le excel de generale
@login_required(login_url='/login/')
@gestion_required
def generate_excel_generale(request):
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    wb = Workbook()
    ws = wb.active

    headers = [
        'Projet', 'Nom de l\'organisation', 'Nature de l\'organisation', 'Sigle', 'Pays d\'origine', 'Région', 'Province', 'Commune', 'Village/Secteur/Avenue/Rue',
        'Boite postale', 'Numéro de téléphone fixe', 'Numéro de téléphone mobile', 'Adresse mail professionnelle', 'Site web', 'Nom et Prénom(s)', 'Nationalité',
        'Fonction(Président,...)', 'Numéro fixe', 'Numéro mobile', 'Dernier renouvèlement des Instances dirigeantes', 'Dernière Assemblée Générale Ordinaire',
        'Dernière session statutaire du bureau exécutif', 'Durée du mandat du bureau exécutif', 'Objectifs', 'Groupes cibles specifique', 'Hommes Employés nationaux Contrat à Durée Indéterminée (CDI)',
        'Femmes Employés nationaux Contrat à Durée Indéterminée (CDI)', 'Hommes Employés nationaux Contrat à Durée déterminée (CDD)', 'Femmes Employés nationaux Contrat à Durée déterminée (CDD)',
        'Hommes Employés expatriés Contrat à Durée Indéterminée (CDI)', 'Femmes Employés expatriés Contrat à Durée Indéterminée (CDI)', 'Hommes Employés expatriés Contrat à Durée déterminée (CDD)',
        'Femmes Employés expatriés Contrat à Durée déterminée (CDD)', 'Hommes Bénévoles ou volontaires Nationaux', 'Femmes Bénévoles ou volontaires Nationaux',
        'Hommes Bénévoles ou volontaires Expatriés', 'Femmes Bénévoles ou volontaires Expatriés', 'Hommes Personnel de l\'Administration publique en détachement', 'Femmes Personnel de l\'Administration publique en détachement',
        'Nom partenariats', 'N° de convention de partenariat / protocole d\'entente', 'Date de début d\'effet', 'Date de fin d\'effet'
    ]
    ws.append(headers)
    

    # remplir les donnees
    for info in infos_test:
        partenariats = partenariats_dict[info.id]
        objectifs = objectifs_dict[info.id]

        objectifs_str = ', '.join([o['objectifs'] for o in objectifs])

        if partenariats:
            for p in partenariats:
                row = [
                    info.id_projet.nom, info.nom_org , info.nature_org, info.sigle, info.pays_origine, info.region, info.province, info.commune,
                    info.village, info.boite_postale, info.numb_fixe, info.numb_mobile, info.adresse_mail, info.site_web, info.nom_complet_resp,
                    info.nationalite_resp, info.fonction_resp, info.numb_fixe_resp, info.numb_mobile_resp, info.renou_instance, info.assem_general,
                    info.session_statut, info.mandat_bureau,  objectifs_str, info.groupes_cibles, info.em_nation_cdi_homme, info.em_nation_cdi_femme,
                    info.em_nation_cdd_homme, info.em_nation_cdd_femme, info.em_expa_cdi_homme, info.em_expa_cdi_femme, info.em_expa_cdd_homme,
                    info.em_expa_cdd_femme, info.benevol_nation_homme, info.benevol_nation_femme, info.benevol_expa_homme, info.benevol_expa_femme,
                    info.personnel_admin_homme, info.personnel_admin_femme, p['nom'], p['numero'], p['date_debut'], p['date_fin']
                ]
                ws.append(row)
        else:
            row = [
                info.id_projet.nom, info.nom_org, info.nature_org, info.sigle, info.pays_origine, info.region, info.province, info.commune,
                    info.village, info.boite_postale, info.numb_fixe, info.numb_mobile, info.adresse_mail, info.site_web, info.nom_complet_resp,
                    info.nationalite_resp, info.fonction_resp, info.numb_fixe_resp, info.numb_mobile_resp, info.renou_instance, info.assem_general,
                    info.session_statut, info.mandat_bureau,  objectifs_str, info.groupes_cibles, info.em_nation_cdi_homme, info.em_nation_cdi_femme,
                    info.em_nation_cdd_homme, info.em_nation_cdd_femme, info.em_expa_cdi_homme, info.em_expa_cdi_femme, info.em_expa_cdd_homme,
                    info.em_expa_cdd_femme, info.benevol_nation_homme, info.benevol_nation_femme, info.benevol_expa_homme, info.benevol_expa_femme,
                    info.personnel_admin_homme, info.personnel_admin_femme, '', '', '', ''
            ]
            ws.append(row)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=synthese_generales.xlsx'

    wb.save(response)

    return response


# Generer le pdf de synthese planfication
@login_required(login_url='/login/')
@gestion_required
def generate_pdf_planification(request):
    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()

    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)
        
        unite_physique = ', '.join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = ', '.join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = ', '.join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = ', '.join(activite_groupe.values_list('responsable', flat=True))
        part_burkina = sum(activite_groupe.values_list('part_burkina', flat=True))
        budget_total = sum(activite_groupe.values_list('budget_total', flat=True))
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activite=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'part_burkina': part_burkina,
            'budget_total': budget_total,
            'partenaires': list(partenaires_groupes.values())
        })

    # Charger le template HTML
    template_path = 'service/invoice_planification.html'
    template = get_template(template_path)
    context = {
        'activites_groupees': activites_groupees
    }
    html = template.render(context)
    
    # Créer un objet HttpResponse avec le type de contenu PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_planification.pdf"'
    
    # Convertir le template HTML en PDF
    pisa_status = pisa.CreatePDF(html, dest=response)
    
    # Si la conversion a réussi, retourner la réponse avec le PDF généré
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    return response

# Generer le word de planification
@login_required(login_url='/login/')
@gestion_required
def generate_word_planification(request):
    activites_groupees = []
    nom_final = []
    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()

    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)
        for activite in activite_groupe:
            projet = activite.id_projet
            nom_final.append(projet.nom)

        
        unite_physique = ', '.join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = ', '.join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = ', '.join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = ', '.join(activite_groupe.values_list('responsable', flat=True))
        part_burkina = sum(activite_groupe.values_list('part_burkina', flat=True))
        budget_total = sum(activite_groupe.values_list('budget_total', flat=True))
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activite=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'projet': nom_final,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'part_burkina': part_burkina,
            'budget_total': budget_total,
            'partenaires': list(partenaires_groupes.values())
        })

    # creer un document word
    doc = Document()

    # ajouter le titre
    doc.add_heading('SYNTHÈSE GLOBALE PLANIFICATION OPERATIONNELLE', level=1)

    for activite in activites_groupees:
        doc.add_heading(' | '.join(f"{i}" for i in nom_final))
        doc.add_heading(f'Titre', level=2)
        doc.add_paragraph(f"{activite['titre']}")
        doc.add_heading(f'Unité physique', level=2)
        doc.add_paragraph(f"{activite['unite_physique']}")
        doc.add_heading(f'Quantité prévue', level=2)
        doc.add_paragraph(f"{activite['quantite_prevue']}")
        doc.add_heading(f'Période prévue', level=2)
        doc.add_paragraph(f"Debut : {activite['periode_prevue_debut']}")
        doc.add_paragraph(f"Fin : {activite['periode_prevue_fin']}")
        doc.add_heading(f'Responsable d\'exécution', level=2)
        doc.add_paragraph(f"{activite['responsable']}")
        doc.add_heading(f'Budget total prévu (FCFA)', level=2)
        doc.add_paragraph(f"{activite['budget_total']}")
        doc.add_heading(f'Part Etat Burkina Faso (FCFA)', level=2)
        doc.add_paragraph(f"{activite['part_burkina']}")
        doc.add_heading(f'Partenaires financiers', level=2)
        for partenaire in activite['partenaires']:
            doc.add_paragraph(f"Nom du partenaire : {partenaire['nom']}")
            doc.add_paragraph(f"Part du partenaire : {partenaire['part']}")
            doc.add_paragraph("----------------------------------------")

        doc.add_paragraph("_________________________________________________________________________________________________________")

    # Créer un flux de mémoire pour stocker le document Word
    output = BytesIO()
    doc.save(output)

    # Réinitialiser le flux à la position de départ
    output.seek(0)

    # Creer une reponse Http avec le contenu du document word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=synthese_planification.docx'
    response.write(output.getvalue())

    return response    

# Generer le excel de planification
@login_required(login_url='/login/')
@gestion_required
def generate_excel_planification(request):
    activites_groupees = []

    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()

    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)
        
        unite_physique = ', '.join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = ', '.join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = ', '.join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = ', '.join(activite_groupe.values_list('responsable', flat=True))
        part_burkina = sum(activite_groupe.values_list('part_burkina', flat=True))
        budget_total = sum(activite_groupe.values_list('budget_total', flat=True))
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activite=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'part_burkina': part_burkina,
            'budget_total': budget_total,
            'partenaires': list(partenaires_groupes.values())
        })

    # Créer un nouveau classeur Excel
    wb = Workbook()
    # Créer une nouvelle feuille dans le classeur
    ws = wb.active
    # Ajouter des en-têtes de colonnes
    ws.append(['Titre', 'Unité Physique', 'Quantité Prévue', 'Date de début de l\'activité', 
               'Date de fin de l\'activité', 'Responsable', 'Part Burkina', 'Budget Total', 
               'Nom du partenaire', 'Part du partenaire'])
    # Ajouter les données à la feuille Excel
    for activite in activites_groupees:
        titre = activite['titre']
        unite_physique = activite['unite_physique']
        quantite_prevue = activite['quantite_prevue']
        periode_prevue_debut = activite['periode_prevue_debut']
        periode_prevue_fin = activite['periode_prevue_fin']
        responsable = activite['responsable']
        part_burkina = activite['part_burkina']
        budget_total = activite['budget_total']
        partenaires = ', '.join([f"{partenaire['nom']} ({partenaire['part']})" for partenaire in activite['partenaires']])
        nom_partenaire = ', '.join([f"{partenaire['nom']}" for partenaire in activite['partenaires']])
        part_partenaire = ', '.join([f"{partenaire['part']}" for partenaire in activite['partenaires']])

        ws.append([titre, unite_physique, quantite_prevue, periode_prevue_debut, 
                   periode_prevue_fin, responsable, part_burkina, budget_total, 
                   nom_partenaire, part_partenaire])
    
    # Créer une réponse HTTP pour le fichier Excel
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=synthese_planification.xlsx'

    # Enregistrer le classeur Excel dans la réponse HTTP
    wb.save(response)

    return response

# Generer pdf suivi
@login_required(login_url='/login/')
@gestion_required
def generate_pdf_suivi(request):
    distinct_titres = ActivitePlu.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = ActivitePlu.objects.filter(titre=titre)
        commune = ', '.join(activite_groupe.values_list('commune', flat=True))
        province = ', '.join(activite_groupe.values_list('province', flat=True))
        region = ', '.join(activite_groupe.values_list('region', flat=True))
        paroisse = ', '.join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = ', '.join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = ', '.join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = ', '.join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = ', '.join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        part_burkina = sum(activite_groupe.values_list('part_burkina', flat=True))
        total_benef_direct = sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activiteplus=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'part_burkina': part_burkina,
            'partenaires': list(partenaires_groupes.values())
        })
    template_path = 'service/invoice_suivi.html'
    template = get_template(template_path)
    context = {
        'activites_groupees': activites_groupees
    }
    html = template.render(context)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_suivi.pdf"'

    pisa_status = pisa.CreatePDF(html, dest=response)

    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    return response

# Generer le word de suivi
@login_required(login_url='/login/')
@gestion_required
def generate_word_suivi(request):
    distinct_titres = ActivitePlu.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    nom_final = []
    for titre in distinct_titres:
        activite_groupe = ActivitePlu.objects.filter(titre=titre)
        for activite in activite_groupe:
            projet = activite.id_projet
            nom_final.append(projet.nom)

        commune = ', '.join(activite_groupe.values_list('commune', flat=True))
        province = ', '.join(activite_groupe.values_list('province', flat=True))
        region = ', '.join(activite_groupe.values_list('region', flat=True))
        paroisse = ', '.join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = ', '.join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = ', '.join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = ', '.join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = ', '.join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        part_burkina = sum(activite_groupe.values_list('part_burkina', flat=True))
        total_benef_direct = sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activiteplus=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'part_burkina': part_burkina,
            'partenaires': list(partenaires_groupes.values())
            })
    
    doc = Document()

    doc.add_heading('SYNTHÈSE GLOBALE SUIVI', level=1)

    for activite in activites_groupees:
        doc.add_heading(' | '.join(f"{i}" for i in nom_final), level=2)

        doc.add_heading('Titre', level=2)
        doc.add_paragraph(f"{activite['titre']}")

        doc.add_heading('Commune', level=2)
        doc.add_paragraph(f"{activite['commune']}")

        doc.add_heading('Province', level=2)
        doc.add_paragraph(f"{activite['province']}")

        doc.add_heading('Région', level=2)
        doc.add_paragraph(f"{activite['region']}")

        doc.add_heading('Paroisse', level=2)
        doc.add_paragraph(f"{activite['paroisse']}")

        doc.add_heading('Unité physique', level=2)
        doc.add_paragraph(f"{activite['unite_physique']}")

        doc.add_heading('Quantité réalisé', level=2)
        doc.add_paragraph(f"{activite['quantite_prevue']}")

        doc.add_heading('Période réalisé', level=2)
        doc.add_paragraph(f"Date de debut de l'activite : {activite['periode_prevue_debut']}")
        doc.add_paragraph(f"Date de debut de l'activite : {activite['periode_prevue_fin']}")

        doc.add_heading('Responsable d\'exécution', level=2)
        doc.add_paragraph(f"{activite['responsable']}")

        doc.add_heading('Coût de realisation', level=2)
        doc.add_paragraph(f"{activite['cout_realisation']}")

        doc.add_heading('Contribution bénéficiaire', level=2)
        doc.add_paragraph(f"{activite['contribution_beneficiaire']}")

        doc.add_heading('Contribution partenaire', level=2)
        doc.add_paragraph(f"{activite['contribution_partenaire']}")

        doc.add_heading('Nombre bénéficiare direct homme', level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_homme']}")

        doc.add_heading('Nombre bénéficiare direct femme', level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_femme']}")

        doc.add_heading('Nombre total de bénéficiaire direct', level=2)
        doc.add_paragraph(f"{activite['total_benef_direct']}")

        doc.add_heading('Part Etat Burkina Faso (FCFA)', level=2)
        doc.add_paragraph(f"{activite['part_burkina']}")

        doc.add_heading('Partenaires financiers', level=2)
        for partenaire in activite['partenaires']:
            doc.add_paragraph(f"Nom du partenaire : {partenaire['nom']}")
            doc.add_paragraph(f"Part du partenaire : {partenaire['part']}")
            doc.add_paragraph("----------------------------------------")
    
    doc.add_paragraph("_________________________________________________________________________________________________________")

    output = BytesIO()
    doc.save(output)

    output.seek(0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=synthese_suivi.docx'
    response.write(output.getvalue())

    return response

# Generer le excel de suivi
@login_required(login_url='/login/')
@gestion_required
def generate_excel_suivi(request):
    distinct_titres = ActivitePlu.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = ActivitePlu.objects.filter(titre=titre)
        commune = ', '.join(activite_groupe.values_list('commune', flat=True))
        province = ', '.join(activite_groupe.values_list('province', flat=True))
        region = ', '.join(activite_groupe.values_list('region', flat=True))
        paroisse = ', '.join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = ', '.join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = ', '.join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = ', '.join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = ', '.join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        part_burkina = sum(activite_groupe.values_list('part_burkina', flat=True))
        total_benef_direct = sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activiteplus=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'part_burkina': part_burkina,
            'partenaires': list(partenaires_groupes.values())
        })
    
    wb = Workbook()
    ws = wb.active

    ws.append(['Titre', 'Commune', 'Province', 'Région', 'Paroisse', 'Unité physique', 'Quantité réalisé', 'Date de début de l\'activité', 'Date de fin de l\'activité', 'Responsable d\'exécution', 'Coût de réalisation', 'Contribution bénéficiaire', 'Contribution partenaire', 'Nombre bénéficiare direct homme', 'Nombre bénéficiare direct femme', 'Nombre total de bénéficiaire direct', 'Part Etat Burkina Faso (FCFA)', 
               'Nom du partenaire', 'Part du partenaire'])
    for activite in activites_groupees:
        titre = activite['titre']
        commune = activite['commune']
        province = activite['province']
        region = activite['region']
        paroisse = activite['paroisse']
        unite_physique = activite['unite_physique']
        quantite_realise = activite['quantite_prevue']
        periode_prevue_debut = activite['periode_prevue_debut']
        periode_prevue_fin = activite['periode_prevue_fin']
        responsable = activite['responsable']
        cout_realisation = activite['cout_realisation']
        contribution_beneficiaire = activite['contribution_beneficiaire']
        contribution_partenaire  = activite['contribution_partenaire']
        total_benef_direct = activite['total_benef_direct']
        nbre_benef_direct_homme = activite['nbre_benef_direct_homme']
        nbre_benef_direct_femme = activite['nbre_benef_direct_femme']
        part_burkina = activite['part_burkina']
        nom_partenaire = ', '.join([f"{partenaire['nom']}" for partenaire in activite['partenaires']])
        part_partenaire = ', '.join([f"{partenaire['part']}" for partenaire in activite['partenaires']])

        ws.append([titre, commune, province, region, paroisse, unite_physique, quantite_realise, periode_prevue_debut, periode_prevue_fin, responsable, cout_realisation, contribution_beneficiaire, contribution_partenaire, nbre_benef_direct_homme, nbre_benef_direct_femme, total_benef_direct, part_burkina, nom_partenaire,
               part_partenaire])
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=synthese_suivi.xlsx'

    wb.save(response)

    return response

@login_required(login_url='/login/')
@gestion_required
def globale_generale(request):
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    context = {
        'infos_test': infos_test,
        'partenariats_dict': partenariats_dict,
        'objectifs_dict': objectifs_dict
    }

    return render(request, 'service/globale_generale.html', context)



@login_required(login_url='/login/')
@gestion_required
def globale_planification(request):
    #
    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)
        
        unite_physique = ', '.join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = ', '.join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = ', '.join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = ', '.join(activite_groupe.values_list('responsable', flat=True))
        part_burkina = sum(activite_groupe.values_list('part_burkina', flat=True))
        budget_total = sum(activite_groupe.values_list('budget_total', flat=True))
        
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activite=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'part_burkina': part_burkina,
            'budget_total': budget_total,
            'partenaires': list(partenaires_groupes.values())
        })
    context = {
        'projets': 'projets',
        'activites_groupees': activites_groupees,
        'activieplus': 'activiteplus'
    }
    return render(request, 'service/globale.html', context)

@login_required(login_url='/login/')
@gestion_required
def globale_suivi(request):
    #
    distinct_titres = ActivitePlu.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = ActivitePlu.objects.filter(titre=titre)
        commune = ', '.join(activite_groupe.values_list('commune', flat=True))
        province = ', '.join(activite_groupe.values_list('province', flat=True))
        region = ', '.join(activite_groupe.values_list('region', flat=True))
        paroisse = ', '.join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = ', '.join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = ', '.join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = ', '.join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = ', '.join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        part_burkina = sum(activite_groupe.values_list('part_burkina', flat=True))
        total_benef_direct = sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activiteplus=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'part_burkina': part_burkina,
            'partenaires': list(partenaires_groupes.values())
        })
    context = {
        'projets': 'projets',
        'activites_groupees': activites_groupees,
        'activieplus': 'activiteplus'
    }
    return render(request, 'service/globale1.html', context)


@login_required(login_url='/login/')
def index(request):
    user = request.user
    total_service = Projet.objects.count()
    total_service_user = Projet.objects.filter(utilisateur=user).count()
    context = {
        'total_service' : total_service,
        'user': user,
        'total_service_user': total_service_user
    }
    
    return render(request, 'base.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def service(request):
    user = request.user
    service = Projet.objects.all()
    service_plus = Projet.objects.filter(utilisateur=user)
    context = {
        'service' : service
    }

    return render(request, 'service/index.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def view_activite(request, projet_id, activite_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = get_object_or_404(Activite, pk=activite_id)
    partenaires = Partenaire.objects.filter(id_activite=activite_id) 
    context = {
        'activite': activite,
        'projet': projet,
        'partenaires': partenaires,
        'user': user
    }

    return render(request, 'service/view_activite.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def view_activiteplus(request, projet_id, activite_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = get_object_or_404(ActivitePlu, pk=activite_id)
    partenaires = Partenaire.objects.filter(id_activiteplus=activite_id) 
    context = {
        'activite': activite,
        'projet': projet,
        'partenaires': partenaires,
        'user': user
    }

    return render(request, 'service/view_activiteplus.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def view_generale(request, projet_id, gen_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    infos = get_object_or_404(InfosGenerale, pk=gen_id)
    objectifs = Objectif.objects.filter(id_general=gen_id)
    partenariats = Partenariat.objects.filter(id_general=gen_id)
    context = {
        'infos': infos,
        'projet': projet,
        'objectifs': objectifs,
        'user': user,
        'partenariats': partenariats
    }

    return render(request, 'service/view_generale.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def profils(request):

    return render(request, 'service/profil.html')

@login_required(login_url='/login/')
@gestion_user_required
def ajouter_planification(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    titres = TitreActivite.objects.all()
    context = {
            'projet': projet,
            'titres': titres,
            'user': user
        }
    if request.method == 'POST':
        titre_nom = request.POST.get('select-activite-titre')
        unite = request.POST.get('unite')
        quantite = request.POST.get('quantite')
        periode_debut = request.POST.get('periode_debut')
        periode_fin = request.POST.get('periode_fin')
        responsable = request.POST.get('responsable')
        budget = request.POST.get('budget_total')
        part_burkina = request.POST.get('part_burkina')
        partenaires = request.POST.getlist('nom_partenaire[]')
        parts = request.POST.getlist('part_partenaire[]')

        if Activite.objects.filter(titre=titre_nom, id_projet=projet).exists():
            messages.info(request, "Cet activite existe deja dans la base")
            return render(request, 'service/ajouter_planification.html', context)
        else:
            if TitreActivite.objects.filter(titre=titre_nom).exists():
                titre_obj = titre_nom
                activite = Activite.objects.create(
                    utilisateur = user,
                    titre = titre_obj,
                    unite_physique = unite,
                    quantite_prevue = quantite,
                    periode_prevue_debut = periode_debut,
                    periode_prevue_fin = periode_fin,
                    responsable = responsable,
                    budget_total = budget,
                    part_burkina = part_burkina,
                    id_projet = projet
                )
                for name, part in zip(partenaires, parts):
                    testing = Partenaire.objects.create(nom=name, part=part, id_activite=activite)
                    testing.save()
                
                messages.success(request, "Activite enregistrer avec succes")
                return render(request, 'service/ajouter_planification.html', context)
            else:
                messages.info(request, 'Cet activite n\'exite pas dans la liste')
                return render(request, 'service/ajouter_planification.html', context)
    else:
        return render(request, 'service/ajouter_planification.html', context)

@login_required(login_url='/login/')
@gestion_required
def ajouter_general(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    titres = TitreActivite.objects.all()
    context = {
            'projet': projet,
            'titres': titres,
            'user': user
        }
    if request.method == 'POST':
        nom_org = request.POST.get('nom_org')
        nature_org = request.POST.get('nature_org')
        sigle = request.POST.get('sigle')
        pays_origine = request.POST.get('pays_ori')
        region = request.POST.get('region')
        province = request.POST.get('province')
        commune = request.POST.get('commune')
        village = request.POST.get('village')
        boite_postale = request.POST.get('boite_postale')
        numb_mobile = request.POST.get('ong_mobile')
        numb_fixe = request.POST.get('ong_fixe')
        adresse_mail = request.POST.get('ong_email_pro')
        site_web = request.POST.get('site_web')
        #
        nom_complet_resp = request.POST.get('resp_nom')
        nationalite_resp = request.POST.get('resp_nationalite')
        fonction_resp = request.POST.get('resp_fonction')
        numb_fixe_resp = request.POST.get('resp_fixe')
        numb_mobile_resp = request.POST.get('resp_mobile')
        #
        renou_instance = request.POST.get('date1')
        assem_general = request.POST.get('date2')
        session_statut = request.POST.get('date3')
        mandat_bureau = request.POST.get('date4')
        #
        nom_complet_canevas = request.POST.get('repondant_nom')
        numb_fixe_canevas = request.POST.get('repondant_fixe')
        numb_mobile_canevas = request.POST.get('repondant_mobile')
        adresse_mail_canevas = request.POST.get('repondant_adresse')
        #
        objectifs = request.POST.getlist('objectifs[]')
        #
        groupes_cibles = request.POST.getlist('groupe_check[]')
        autre_groupe = request.POST.get('autre_groupe')
        #
        em_nation_cdi_homme = request.POST.get('em_nation_cdi_homme')
        em_nation_cdi_femme = request.POST.get('em_nation_cdi_femme')
        #
        em_nation_cdd_homme = request.POST.get('em_nation_cdd_homme')
        em_nation_cdd_femme = request.POST.get('em_nation_cdd_femme')
        #
        em_expa_cdi_homme = request.POST.get('em_expa_cdi_homme')
        em_expa_cdi_femme = request.POST.get('em_expa_cdi_femme')
        #
        em_expa_cdd_homme = request.POST.get('em_expa_cdd_homme')
        em_expa_cdd_femme = request.POST.get('em_expa_cdd_femme')
        #
        benevol_nation_homme = request.POST.get('benevol_nation_homme')
        benevol_nation_femme = request.POST.get('benevol_nation_femme')
        #
        benevol_expa_homme = request.POST.get('benevol_expa_homme')
        benevol_expa_femme = request.POST.get('benevol_expa_femme')
        #
        personnel_admin_homme = request.POST.get('personnel_admin_homme')
        personnel_admin_femme = request.POST.get('personnel_admin_femme')
        # ajouter un ou plusieurs partenaires
        nom_partenariat = request.POST.getlist('partenariat[]')
        numero_partenariat = request.POST.getlist('number_part[]')
        date_debut_part = request.POST.getlist('date_debut_part[]')
        date_fin_part = request.POST.getlist('date_fin_part[]')

        if InfosGenerale.objects.filter(id_projet=projet).exists():
            messages.info(request, 'Ce projet a deja des informations generales')
            return render(request, 'service/choix_projet.html', context)
        else:
            if all([
                nom_org,
                nature_org,
                sigle,
                pays_origine,
                region,
                province,
                commune,
                village,
                boite_postale,
                numb_fixe,
                numb_mobile,
                adresse_mail,
                site_web,
                nom_complet_resp,
                nationalite_resp,
                fonction_resp,
                numb_fixe_resp,
                numb_mobile_resp,
                renou_instance,
                assem_general,
                session_statut,
                mandat_bureau,
                nom_complet_canevas,
                numb_fixe_canevas,
                numb_mobile_canevas,
                adresse_mail_canevas,
                groupes_cibles,
                em_nation_cdi_homme,
                em_nation_cdi_femme,
                em_nation_cdd_homme,
                em_nation_cdd_femme,
                em_expa_cdi_homme,
                em_expa_cdi_femme,
                em_expa_cdd_homme,
                em_expa_cdd_femme,
                benevol_nation_homme,
                benevol_nation_femme,
                benevol_expa_homme,
                benevol_expa_femme,
                personnel_admin_homme,
                personnel_admin_femme
            ]):
                infos = InfosGenerale.objects.create(
                    utilisateur = user,
                    id_projet = projet,
                    nom_org = nom_org,
                    nature_org = nature_org,
                    sigle = sigle,
                    pays_origine = pays_origine,
                    region = region,
                    province = province,
                    commune = commune,
                    village = village,
                    boite_postale = boite_postale,
                    numb_mobile = numb_mobile,
                    numb_fixe = numb_fixe,
                    adresse_mail = adresse_mail,
                    site_web = site_web,
                    nom_complet_resp = nom_complet_resp,
                    nationalite_resp = nationalite_resp,
                    fonction_resp = fonction_resp,
                    numb_fixe_resp = numb_fixe_resp,
                    numb_mobile_resp = numb_mobile_resp,
                    renou_instance = renou_instance,
                    assem_general = assem_general,
                    session_statut = session_statut,
                    mandat_bureau = mandat_bureau,
                    nom_complet_canevas = nom_complet_canevas,
                    numb_fixe_canevas = numb_fixe_canevas,
                    numb_mobile_canevas = numb_mobile_canevas,
                    adresse_mail_canevas = adresse_mail_canevas,
                    groupes_cibles = groupes_cibles,
                    em_nation_cdi_homme = em_nation_cdi_homme,
                    em_nation_cdi_femme = em_nation_cdi_femme,
                    em_nation_cdd_homme = em_nation_cdd_homme,
                    em_nation_cdd_femme = em_nation_cdd_femme,
                    em_expa_cdi_homme = em_expa_cdi_homme,
                    em_expa_cdi_femme = em_expa_cdi_femme,
                    em_expa_cdd_homme = em_expa_cdd_homme,
                    em_expa_cdd_femme = em_expa_cdd_femme,
                    benevol_nation_homme = benevol_nation_homme,
                    benevol_nation_femme = benevol_nation_femme,
                    benevol_expa_homme = benevol_expa_homme,
                    benevol_expa_femme = benevol_expa_femme,
                    personnel_admin_homme = personnel_admin_homme,
                    personnel_admin_femme = personnel_admin_femme,
                    autre_groupe = autre_groupe
                )
                for obj in objectifs:
                    objectif = Objectif.objects.create(objectifs=obj, id_general=infos)
                    objectif.save()

                for name, numero, date_debut, date_fin in zip(nom_partenariat, numero_partenariat, date_debut_part, date_fin_part):
                    partenariat = Partenariat.objects.create(
                        nom = name,
                        numero = numero,
                        date_debut = date_debut,
                        date_fin = date_fin,
                        id_general = infos
                    )
                    partenariat.save()


                messages.success(request, 'Les informations generales ont ete creer avec success')
                return render(request, 'service/choix_projet.html', context)
            else:
                messages.info(request, 'Veuillez renseigner tout les informations')
                return render(request, 'service/ajouter_general.html', context)
    else:
        return render(request, 'service/ajouter_general.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def ajouter_suivi(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    titres = TitreActivite.objects.all()
    context = {
            'projet': projet,
            'titres': titres,
            'user': user
        }
    if request.method == 'POST':
        titre_nom = request.POST.get('select-activite-titre')
        commune = request.POST.get('commune')
        province = request.POST.get('province')
        region = request.POST.get('region')
        paroisse = request.POST.get('paroisse')
        unite = request.POST.get('unite')
        quantite = request.POST.get('quantite')
        periode_debut = request.POST.get('periode_debut')
        periode_fin = request.POST.get('periode_fin')
        responsable = request.POST.get('responsable')
        cout_realisation = request.POST.get('cout_realisation')
        contribution_beneficiaire = request.POST.get('contrib_benef')
        contribution_partenaire = request.POST.get('contrib_part')
        total_benef_direct = request.POST.get('total_benef')
        nbre_benef_direct_homme = request.POST.get('nbre_ben_homme')
        nbre_benef_direct_femme = request.POST.get('nbre_ben_femme')
        part_burkina = request.POST.get('part_burkina')
        partenaires = request.POST.getlist('nom_partenaire[]')
        parts = request.POST.getlist('part_partenaire[]')
        
        if ActivitePlu.objects.filter(titre=titre_nom, id_projet=projet).exists():
            messages.info(request, "Cet activite existe deja dans la base")
            return render(request, 'service/ajouter_suivi.html', context)
        else:
            if TitreActivite.objects.filter(titre=titre_nom).exists():
                titre_obj = titre_nom
                activite = ActivitePlu.objects.create(
                    utilisateur = user,
                    titre = titre_obj,
                    commune = commune,
                    province = province,
                    region = region,
                    paroisse = paroisse,
                    cout_realisation = cout_realisation,
                    contribution_beneficiaire = contribution_beneficiaire,
                    contribution_partenaire = contribution_partenaire,
                    total_benef_direct = total_benef_direct,
                    nbre_benef_direct_homme = nbre_benef_direct_homme,
                    nbre_benef_direct_femme = nbre_benef_direct_femme,
                    unite_physique = unite,
                    quantite_prevue = quantite,
                    periode_prevue_debut = periode_debut,
                    periode_prevue_fin = periode_fin,
                    responsable = responsable,
                    part_burkina = part_burkina,
                    id_projet = projet
                )
                for name, part in zip(partenaires, parts):
                    partenaire = Partenaire.objects.create(nom=name, part=part, id_activiteplus=activite)
                
                messages.success(request, "Activite enregistrer avec succes")
                return render(request, 'service/ajouter_suivi.html', context)
            else:
                messages.info(request, 'Cet activite n\'exite pas dans la liste')
                return render(request, 'service/ajouter_suivi.html', context)
    else:
        return render(request, 'service/ajouter_suivi.html', context)

@login_required(login_url='/login/')
@gestion_required
def modifier_general(request, projet_id, gen_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    infos = get_object_or_404(InfosGenerale, id=gen_id)
    objectifs = Objectif.objects.filter(id_general=gen_id)
    partenariats = Partenariat.objects.filter(id_general=gen_id)
    context = {
            'projet': projet,
            'infos': infos,
            'objectifs': objectifs,
            'user': user,
            'partenariats': partenariats
        }
    
    if request.method == 'POST':
        nom_org = request.POST.get('nom_org')
        nature_org = request.POST.get('nature_org')
        sigle = request.POST.get('sigle')
        pays_origine = request.POST.get('pays_ori')
        region = request.POST.get('region')
        province = request.POST.get('province')
        commune = request.POST.get('commune')
        village = request.POST.get('village')
        boite_postale = request.POST.get('boite_postale')
        numb_mobile = request.POST.get('ong_mobile')
        numb_fixe = request.POST.get('ong_fixe')
        adresse_mail = request.POST.get('ong_email_pro')
        site_web = request.POST.get('site_web')
        #
        nom_complet_resp = request.POST.get('resp_nom')
        nationalite_resp = request.POST.get('resp_nationalite')
        fonction_resp = request.POST.get('resp_fonction')
        numb_fixe_resp = request.POST.get('resp_fixe')
        numb_mobile_resp = request.POST.get('resp_mobile')
        #
        renou_instance = request.POST.get('date1')
        assem_general = request.POST.get('date2')
        session_statut = request.POST.get('date3')
        mandat_bureau = request.POST.get('date4')
        #
        nom_complet_canevas = request.POST.get('repondant_nom')
        numb_fixe_canevas = request.POST.get('repondant_fixe')
        numb_mobile_canevas = request.POST.get('repondant_mobile')
        adresse_mail_canevas = request.POST.get('repondant_adresse')
        #
        objectifs = request.POST.getlist('objectifs[]')
        #
        groupes_cibles = request.POST.getlist('groupe_check[]')
        #
        em_nation_cdi_homme = request.POST.get('em_nation_cdi_homme')
        em_nation_cdi_femme = request.POST.get('em_nation_cdi_femme')
        #
        em_nation_cdd_homme = request.POST.get('em_nation_cdd_homme')
        em_nation_cdd_femme = request.POST.get('em_nation_cdd_femme')
        #
        em_expa_cdi_homme = request.POST.get('em_expa_cdi_homme')
        em_expa_cdi_femme = request.POST.get('em_expa_cdi_femme')
        #
        em_expa_cdd_homme = request.POST.get('em_expa_cdd_homme')
        em_expa_cdd_femme = request.POST.get('em_expa_cdd_femme')
        #
        benevol_nation_homme = request.POST.get('benevol_nation_homme')
        benevol_nation_femme = request.POST.get('benevol_nation_femme')
        #
        benevol_expa_homme = request.POST.get('benevol_expa_homme')
        benevol_expa_femme = request.POST.get('benevol_expa_femme')
        #
        personnel_admin_homme = request.POST.get('personnel_admin_homme')
        personnel_admin_femme = request.POST.get('personnel_admin_femme')
        # ajouter un ou plusieurs partenaires
        nom_partenariat = request.POST.getlist('partenariat[]')
        numero_partenariat = request.POST.getlist('number_part[]')
        date_debut_part = request.POST.getlist('date_debut_part[]')
        date_fin_part = request.POST.getlist('date_fin_part[]')
        

        if all([
                nom_org,
                nature_org,
                sigle,
                pays_origine,
                region,
                province,
                commune,
                village,
                boite_postale,
                numb_fixe,
                numb_mobile,
                adresse_mail,
                site_web,
                nom_complet_resp,
                nationalite_resp,
                fonction_resp,
                numb_fixe_resp,
                numb_mobile_resp,
                renou_instance,
                assem_general,
                session_statut,
                mandat_bureau,
                nom_complet_canevas,
                numb_fixe_canevas,
                numb_mobile_canevas,
                adresse_mail_canevas,
                groupes_cibles,
                em_nation_cdi_homme,
                em_nation_cdi_femme,
                em_nation_cdd_homme,
                em_nation_cdd_femme,
                em_expa_cdi_homme,
                em_expa_cdi_femme,
                em_expa_cdd_homme,
                em_expa_cdd_femme,
                benevol_nation_homme,
                benevol_nation_femme,
                benevol_expa_homme,
                benevol_expa_femme,
                personnel_admin_homme,
                personnel_admin_femme
            ]):
            with transaction.atomic():
                generales = InfosGenerale.objects.get(id=gen_id)
                # mettre a jour les champs
                generales.nom_org = nom_org
                generales.nature_org = nature_org
                generales.sigle = sigle
                generales.pays_origine = pays_origine
                generales.region = region
                generales.province = province
                generales.commune = commune
                generales.village = village
                generales.boite_postale = boite_postale
                generales.numb_mobile = numb_mobile
                generales.numb_fixe = numb_fixe
                generales.adresse_mail = adresse_mail
                generales.site_web = site_web
                generales.nom_complet_resp = nom_complet_resp
                generales.nationalite_resp = nationalite_resp
                generales.fonction_resp = fonction_resp
                generales.numb_fixe_resp = numb_fixe_resp
                generales.numb_mobile_resp = numb_mobile_resp
                generales.renou_instance = renou_instance
                generales.assem_general = assem_general
                generales.session_statut = session_statut
                generales.mandat_bureau = mandat_bureau
                generales.nom_complet_canevas = nom_complet_canevas
                generales.numb_fixe_canevas = numb_fixe_canevas
                generales.numb_mobile_canevas = numb_mobile_canevas
                generales.adresse_mail_canevas = adresse_mail_canevas
                generales.groupes_cibles = groupes_cibles
                generales.em_nation_cdi_homme = em_nation_cdi_homme
                generales.em_nation_cdi_femme = em_nation_cdi_femme
                generales.em_nation_cdd_homme = em_nation_cdd_homme
                generales.em_nation_cdd_femme = em_nation_cdd_femme
                generales.em_expa_cdi_homme = em_expa_cdi_homme
                generales.em_expa_cdi_femme = em_expa_cdi_femme
                generales.em_expa_cdd_homme = em_expa_cdd_homme
                generales.em_expa_cdd_femme = em_expa_cdd_femme
                generales.benevol_nation_homme = benevol_nation_homme
                generales.benevol_nation_femme = benevol_nation_femme
                generales.benevol_expa_homme = benevol_expa_homme
                generales.benevol_expa_femme = benevol_expa_femme
                generales.personnel_admin_homme = personnel_admin_homme
                generales.personnel_admin_femme = personnel_admin_femme
                # sauvegarder
                generales.save()

                # objectifs
                for obj in objectifs:
                    try:
                        Objectif.objects.get(objectifs=obj, id_general=generales)
                    except ObjectDoesNotExist:
                        Objectif.objects.create(objectifs=obj, id_general=generales)
                
                # partenariats
                for name, numero, date_debut, date_fin in zip(nom_partenariat, numero_partenariat, date_debut_part, date_fin_part):
                    try:
                        Partenariat.objects.get(
                            nom = name,
                            numero = numero,
                            date_debut = date_debut,
                            date_fin = date_fin,
                            id_general = generales
                        )
                    except ObjectDoesNotExist:
                        partenariat = Partenariat.objects.create(
                            nom = name,
                            numero = numero,
                            date_debut = date_debut,
                            date_fin = date_fin,
                            id_general = generales
                        )

        messages.success(request, 'Les informations generales ont ete modifies avec success')
        return render(request, 'service/choix_projet.html', context)

    else:
        return render(request, 'service/modifier_general.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def modifier_planification(request, projet_id, activite_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = get_object_or_404(Activite, pk=activite_id)
    partenaires = Partenaire.objects.filter(id_activite=activite_id) 
    context = {
        'activite': activite,
        'projet': projet,
        'partenaires': partenaires,
        'user': user
    }
    if request.method == 'POST':
        unite = request.POST.get('unite')
        quantite = request.POST.get('quantite')
        periode_debut = request.POST.get('periode_debut')
        periode_fin = request.POST.get('periode_fin')
        responsable = request.POST.get('responsable')
        budget = request.POST.get('budget_total')
        part_burkina = request.POST.get('part_burkina')
        # liste des partenaires
        partenaires = request.POST.getlist('nom_partenaire[]')
        parts = request.POST.getlist('part_partenaire[]')

        if all([unite, quantite, periode_debut, periode_fin, responsable, budget, part_burkina, partenaires, parts]):
            with transaction.atomic():
                activite = Activite.objects.get(id=activite_id)

                # Mettre a jours les champs de l'activite
                activite.unite_physique = unite
                activite.quantite_prevue = quantite
                activite.periode_prevue_debut = periode_debut
                activite.periode_prevue_fin = periode_fin
                activite.responsable = responsable
                activite.budget_total = budget
                activite.part_burkina = part_burkina
                # sauvegarder les modifications
                activite.save()
            
            # Creer les partenaires associes a l'activite
                for name, part in zip(partenaires, parts):
                    try:
                        # verifier si ce partenaire existe deja
                        Partenaire.objects.get(nom=name, part=part, id_activite=activite)
                        # s'il existe ne rien faire
                    except ObjectDoesNotExist:
                        partenaire = Partenaire.objects.create(nom = name, part = part, id_activite = activite)
            
        messages.success(request, 'Cet activite a ete modifier avec success')
        return render(request, 'service/choix_projet.html', context)
    else:
        return render(request, 'service/modifier_planification.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def modifier_suivi(request, projet_id, activite_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = get_object_or_404(ActivitePlu, pk=activite_id)
    partenaires = Partenaire.objects.filter(id_activite=activite_id) 
    context = {
        'activite': activite,
        'projet': projet,
        'partenaires': partenaires,
        'user': user
    }
    if request.method == 'POST':
        commune = request.POST.get('commune')
        province = request.POST.get('province')
        region = request.POST.get('region')
        paroisse = request.POST.get('paroisse')
        unite = request.POST.get('unite')
        quantite = request.POST.get('quantite')
        periode_debut = request.POST.get('periode_debut')
        periode_fin = request.POST.get('periode_fin')
        responsable = request.POST.get('responsable')
        cout_realisation = request.POST.get('cout_realisation')
        contribution_beneficiaire = request.POST.get('contrib_benef')
        contribution_partenaire = request.POST.get('contrib_part')
        total_benef_direct = request.POST.get('total_benef')
        nbre_benef_direct_homme = request.POST.get('nbre_ben_homme')
        nbre_benef_direct_femme = request.POST.get('nbre_ben_femme')
        part_burkina = request.POST.get('part_burkina')
        # Partenaires financiers
        partenaire = request.POST.getlist('nom_partenaire[]')
        parts = request.POST.getlist('part_partenaire[]')

        if all([commune, province, region, paroisse, unite, quantite, periode_debut, periode_fin, responsable, cout_realisation, contribution_beneficiaire, contribution_partenaire, total_benef_direct, nbre_benef_direct_homme, nbre_benef_direct_femme, parts, part_burkina, partenaire]):
            with transaction.atomic():
                activite = ActivitePlu.objects.get(id=activite_id)
                # Mettre a jour les champs
                activite.commune = commune
                activite.province = province
                activite.region = region
                activite.paroisse = paroisse
                activite.unite_physique = unite
                activite.quantite_prevue = quantite
                activite.periode_prevue_debut = periode_debut
                activite.periode_prevue_fin = periode_fin
                activite.responsable = responsable
                activite.cout_realisation = cout_realisation
                activite.contribution_beneficiaire = contribution_beneficiaire
                activite.contribution_partenaire = contribution_partenaire
                activite.part_burkina = part_burkina
                activite.nbre_benef_direct_homme = nbre_benef_direct_homme
                activite.nbre_benef_direct_femme = nbre_benef_direct_femme
                activite.total_benef_direct = total_benef_direct

                # sauvegarder
                activite.save()

                # modifier les partenaires associes
                for name, part in zip(partenaires, parts):
                    try:
                        Partenaire.objects.get(nom=name, part=part, id_activite=activite)
                    except ObjectDoesNotExist:
                        partenaire = Partenaire.objects.create(nom=name, part=part, id_activite=activite)

        messages.success(request, 'Cet activite a ete modifier avec success')
        return render(request, 'service/choix_projet.html', context)
    else:
        return render(request, 'service/modifier_suivi.html', context)


@login_required(login_url='/login/')
@gestion_user_required
def create_projet(request):
    user = request.user
    context = {
        'user': user
    }
    if request.method == 'POST':
        total_projet = Projet.objects.filter(utilisateur=user).count()
        new_numb = total_projet + 1
        new_name = f"Projet {new_numb}"

        project = Projet.objects.create(nom=new_name, utilisateur=user)

        messages.success(request, 'Projet creer avec success')
        return redirect('projet')
    #render(request, 'service/projet.html', context)
    else:
        return render(request, 'service/create_projet.html', context)


@login_required(login_url='/login/')
@gestion_user_required 
def success(request):
    return render(request, 'service/success.html')

@login_required(login_url='/login/')
@gestion_user_required
def choisir_ajouter(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)

    context = {
        'projet': projet,
        'user': user
    }
    return render(request, 'service/choisir_ajouter.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def choisir_modifier(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)

    context = {
        'projet': projet,
        'user': user
    }
    return render(request, 'service/choisir_modifier.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def projet(request):
    user = request.user
    projects = Projet.objects.filter(utilisateur=user)
    project_users = Projet.objects.filter(utilisateur=user)
    context = {
        'projects': projects,
        'user': user,
        'project_user': project_users
    }
    return render(request, 'service/projet.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def activites(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    # filtrer les activites suivant les users
    activites = Activite.objects.filter(utilisateur=user)

    context = {
        'projet': projet,
        'activites': activites,
        'user': user
    }
    return render(request, 'service/activites.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def activitesplus(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activites = ActivitePlu.objects.filter(utilisateur=user)

    context = {
        'projet': projet,
        'activites': activites,
        'user': user
    }
    return render(request, 'service/activitesplus.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def choix_projet(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = Activite.objects.all()

    context = {
        'projet': projet,
        'activite': activite,
        'user': user
    }

    return render(request, 'service/choix_projet.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def generales(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    infos_generale = InfosGenerale.objects.all()
    infos_generale_users = InfosGenerale.objects.filter(utilisateur=user)

    context = {
        'projet': projet,
        'infos_generale': infos_generale,
        'user': user,
        'infos_generale_users': infos_generale_users
    }

    return render(request, 'service/generales.html', context)

@login_required(login_url='/login/')
@gestion_user_required
def change_password(request):
    if request.method == 'POST':
        old_pass = request.POST.get('old_password')
        new_pass1 = request.POST.get('new_password1')
        new_pass2 = request.POST.get('new_password2')

        #  verifier l'ancien mdp
        if not request.user.check_password(old_pass):
            messages.error(request, 'Votre ancien mot de passe est incorrect')
            return redirect('profils')
        
        # verifier les deux nouveaux mdp
        if new_pass1 != new_pass2:
            messages.error(request, 'Les nouveaux mots de passe ne correspondent pas')
            return redirect('profils')
        
        request.user.set_password(new_pass1)
        request.user.save()
        update_session_auth_hash(request, request.user)
        messages.success(request, 'Votre mot de passe a été mis à jour avec succès.')
        return redirect('profils')
        
    return  render(request, 'service/profil.html')
